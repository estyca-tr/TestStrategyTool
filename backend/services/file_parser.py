"""
File parser service for extracting text from uploaded documents.
Supports PDF, DOCX, and Markdown files.
"""

import os
from typing import Optional


def extract_text_from_file(file_path: str, file_ext: str) -> Optional[str]:
    """
    Extract text content from a file based on its extension.
    
    Args:
        file_path: Path to the file
        file_ext: File extension (e.g., '.pdf', '.docx')
    
    Returns:
        Extracted text content or None if extraction fails
    """
    try:
        if file_ext == '.pdf':
            return extract_from_pdf(file_path)
        elif file_ext in ['.docx', '.doc']:
            return extract_from_docx(file_path)
        elif file_ext == '.md':
            return extract_from_markdown(file_path)
        elif file_ext == '.txt':
            return extract_from_text(file_path)
        else:
            return None
    except Exception as e:
        print(f"Error extracting text from {file_path}: {e}")
        return None


def extract_from_pdf(file_path: str) -> Optional[str]:
    """Extract text from PDF file"""
    try:
        from PyPDF2 import PdfReader
        
        reader = PdfReader(file_path)
        text_parts = []
        
        for page in reader.pages:
            text = page.extract_text()
            if text:
                text_parts.append(text)
        
        return "\n\n".join(text_parts) if text_parts else None
    except Exception as e:
        print(f"PDF extraction error: {e}")
        return None


def extract_from_docx(file_path: str) -> Optional[str]:
    """Extract text from Word document"""
    try:
        from docx import Document
        
        doc = Document(file_path)
        text_parts = []
        
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_parts.append(paragraph.text)
        
        # Also extract from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if row_text:
                    text_parts.append(" | ".join(row_text))
        
        return "\n\n".join(text_parts) if text_parts else None
    except Exception as e:
        print(f"DOCX extraction error: {e}")
        return None


def extract_from_markdown(file_path: str) -> Optional[str]:
    """Extract text from Markdown file"""
    try:
        with open(file_path, 'r', encoding='utf-8') as f:
            return f.read()
    except Exception as e:
        print(f"Markdown extraction error: {e}")
        return None


def extract_from_text(file_path: str) -> Optional[str]:
    """Extract text from plain text file"""
    try:
        with open(file_path, 'r', encoding='utf-8') as f:
            return f.read()
    except Exception as e:
        print(f"Text extraction error: {e}")
        return None

